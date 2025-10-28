#!/usr/bin/env python3
"""
WHY: Microsoft Office and LibreOffice document parsing.
RESPONSIBILITY: Extract text from Word, Excel, ODT, and ODS files.
PATTERNS: Strategy pattern for different office formats, Guard clauses for library checks.
"""

from typing import List, Optional
from artemis_exceptions import DocumentReadError


class WordParser:
    """
    WHY: Parse Microsoft Word documents (.docx, .doc).
    RESPONSIBILITY: Extract text from paragraphs and tables.
    """

    def __init__(self, verbose: bool = False):
        """
        WHY: Initialize parser with library detection.
        RESPONSIBILITY: Check if python-docx is available.

        Args:
            verbose: Enable verbose logging
        """
        self.verbose = verbose
        self.has_docx = self._detect_library()

    def _detect_library(self) -> bool:
        """
        WHY: Check if python-docx library is available.
        RESPONSIBILITY: Return True if docx can be imported.

        Returns:
            True if python-docx is available
        """
        try:
            import docx
            return True
        except ImportError:
            if self.verbose:
                print("[WordParser] Word document support not available. Install: pip install python-docx")
            return False

    def is_available(self) -> bool:
        """Check if Word parsing is available"""
        return self.has_docx

    def parse(self, file_path: str) -> str:
        """
        WHY: Extract text from Word document.
        RESPONSIBILITY: Read paragraphs and tables from .docx file.

        Args:
            file_path: Path to Word document

        Returns:
            Extracted text content

        Raises:
            DocumentReadError: If library not available or parsing fails
        """
        # Guard: Check library availability
        if not self.has_docx:
            raise DocumentReadError(
                "Word document support not available. Install: pip install python-docx",
                context={"file_path": file_path}
            )

        import docx

        try:
            doc = docx.Document(file_path)
            text_content: List[str] = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)

            return "\n".join(text_content)

        except Exception as e:
            raise DocumentReadError(
                f"Error parsing Word document: {str(e)}",
                context={"file_path": file_path},
                original_exception=e
            )


class ExcelParser:
    """
    WHY: Parse Microsoft Excel spreadsheets (.xlsx, .xls).
    RESPONSIBILITY: Extract text from all sheets and cells.
    """

    def __init__(self, verbose: bool = False):
        """
        WHY: Initialize parser with library detection.
        RESPONSIBILITY: Check if openpyxl is available.

        Args:
            verbose: Enable verbose logging
        """
        self.verbose = verbose
        self.has_openpyxl = self._detect_library()

    def _detect_library(self) -> bool:
        """
        WHY: Check if openpyxl library is available.
        RESPONSIBILITY: Return True if openpyxl can be imported.

        Returns:
            True if openpyxl is available
        """
        try:
            import openpyxl
            return True
        except ImportError:
            if self.verbose:
                print("[ExcelParser] Excel support not available. Install: pip install openpyxl")
            return False

    def is_available(self) -> bool:
        """Check if Excel parsing is available"""
        return self.has_openpyxl

    def parse(self, file_path: str) -> str:
        """
        WHY: Extract text from Excel spreadsheet.
        RESPONSIBILITY: Read all sheets and format cells as text.

        Args:
            file_path: Path to Excel file

        Returns:
            Extracted text content with sheet separators

        Raises:
            DocumentReadError: If library not available or parsing fails
        """
        # Guard: Check library availability
        if not self.has_openpyxl:
            raise DocumentReadError(
                "Excel support not available. Install: pip install openpyxl",
                context={"file_path": file_path}
            )

        import openpyxl

        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content: List[str] = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"\n=== Sheet: {sheet_name} ===\n")

                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:
                        text_content.append(" | ".join(row_values))

            return "\n".join(text_content)

        except Exception as e:
            raise DocumentReadError(
                f"Error parsing Excel file: {str(e)}",
                context={"file_path": file_path},
                original_exception=e
            )


class ODTParser:
    """
    WHY: Parse LibreOffice Writer documents (.odt).
    RESPONSIBILITY: Extract text from ODF text documents.
    """

    def __init__(self, verbose: bool = False):
        """
        WHY: Initialize parser with library detection.
        RESPONSIBILITY: Check if odfpy is available.

        Args:
            verbose: Enable verbose logging
        """
        self.verbose = verbose
        self.has_odf = self._detect_library()

    def _detect_library(self) -> bool:
        """
        WHY: Check if odfpy library is available.
        RESPONSIBILITY: Return True if odf can be imported.

        Returns:
            True if odfpy is available
        """
        try:
            from odf import text, teletype
            from odf.opendocument import load
            return True
        except ImportError:
            if self.verbose:
                print("[ODTParser] LibreOffice support not available. Install: pip install odfpy")
            return False

    def is_available(self) -> bool:
        """Check if ODT parsing is available"""
        return self.has_odf

    def parse(self, file_path: str) -> str:
        """
        WHY: Extract text from ODT file.
        RESPONSIBILITY: Read all text paragraphs from ODF document.

        Args:
            file_path: Path to ODT file

        Returns:
            Extracted text content

        Raises:
            DocumentReadError: If library not available or parsing fails
        """
        # Guard: Check library availability
        if not self.has_odf:
            raise DocumentReadError(
                "LibreOffice support not available. Install: pip install odfpy",
                context={"file_path": file_path}
            )

        from odf import text, teletype
        from odf.opendocument import load

        try:
            doc = load(file_path)
            text_content: List[str] = []

            # Extract all text elements
            for element in doc.getElementsByType(text.P):
                paragraph_text = teletype.extractText(element)
                if paragraph_text.strip():
                    text_content.append(paragraph_text)

            return "\n".join(text_content)

        except Exception as e:
            raise DocumentReadError(
                f"Error parsing ODT file: {str(e)}",
                context={"file_path": file_path},
                original_exception=e
            )


class ODSParser:
    """
    WHY: Parse LibreOffice Calc spreadsheets (.ods).
    RESPONSIBILITY: Extract text from ODF spreadsheet tables.
    """

    def __init__(self, verbose: bool = False):
        """
        WHY: Initialize parser with library detection.
        RESPONSIBILITY: Check if odfpy is available.

        Args:
            verbose: Enable verbose logging
        """
        self.verbose = verbose
        self.has_odf = self._detect_library()

    def _detect_library(self) -> bool:
        """
        WHY: Check if odfpy library is available.
        RESPONSIBILITY: Return True if odf can be imported.

        Returns:
            True if odfpy is available
        """
        try:
            from odf.opendocument import load
            from odf.table import Table, TableRow, TableCell
            from odf import teletype
            return True
        except ImportError:
            if self.verbose:
                print("[ODSParser] LibreOffice support not available. Install: pip install odfpy")
            return False

    def is_available(self) -> bool:
        """Check if ODS parsing is available"""
        return self.has_odf

    def parse(self, file_path: str) -> str:
        """
        WHY: Extract text from ODS file.
        RESPONSIBILITY: Read all tables and cells from spreadsheet.

        Args:
            file_path: Path to ODS file

        Returns:
            Extracted text content with sheet separators

        Raises:
            DocumentReadError: If library not available or parsing fails
        """
        # Guard: Check library availability
        if not self.has_odf:
            raise DocumentReadError(
                "LibreOffice support not available. Install: pip install odfpy",
                context={"file_path": file_path}
            )

        from odf.opendocument import load
        from odf.table import Table, TableRow, TableCell
        from odf import teletype

        try:
            doc = load(file_path)
            text_content: List[str] = []

            # Extract all tables
            for table in doc.spreadsheet.getElementsByType(Table):
                table_name = table.getAttribute('name')
                text_content.append(f"\n=== Sheet: {table_name} ===\n")

                for row in table.getElementsByType(TableRow):
                    row_values: List[str] = []
                    for cell in row.getElementsByType(TableCell):
                        cell_text = teletype.extractText(cell)
                        if cell_text:
                            row_values.append(cell_text)

                    if row_values:
                        text_content.append(" | ".join(row_values))

            return "\n".join(text_content)

        except Exception as e:
            raise DocumentReadError(
                f"Error parsing ODS file: {str(e)}",
                context={"file_path": file_path},
                original_exception=e
            )
